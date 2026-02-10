from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
import re
from typing import List, Dict, Tuple
import time
from datetime import datetime, timezone
import io
from .text_utils import clean_text, split_text_into_chunks, normalize_chunk


def process_table(table) -> str:
    """
    Process a table and return its content in a block (Markdown-like) format with clear start/end markers.
    """
    print(f"Processing table with {len(table.rows)} rows...")
    table_lines = []
    # Extract header row
    header_cells = [clean_text(cell.text) for cell in table.rows[0].cells]
    table_lines.append('| ' + ' | '.join(header_cells) + ' |')
    table_lines.append('|' + '|'.join(['---'] * len(header_cells)) + '|')
    # Extract data rows
    for row in table.rows[1:]:
        row_cells = [clean_text(cell.text) for cell in row.cells]
        table_lines.append('| ' + ' | '.join(row_cells) + ' |')
    # Wrap with markers, but only one Table Content:
    return f"Table Content:\n" + '\n'.join(table_lines) + "\nEnd of table content.\n"

def extract_text_from_docx(file) -> List[Tuple[str, str]]:
    """
    Extract text from a Word document, including tables but excluding images.
    Returns a list of tuples (content_type, content) to preserve order.
    """
    try:
        if isinstance(file["content"], bytes):
            file["content"] = Document(io.BytesIO(file["content"]))
        doc = file["content"]
        content_parts = []
        
        # print("\nProcessing document in order...")
        current_paragraphs = []
        
        # Process the document in order
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                para = Paragraph(element, doc)
                # Skip paragraphs that only contain images
                if not any(run._element.xpath('.//w:drawing') for run in para.runs):
                    if para.text.strip():
                        current_paragraphs.append(para.text)
            
            elif element.tag.endswith('tbl'):  # Table
                # Add any accumulated paragraphs first
                if current_paragraphs:
                    content_parts.append(('text', '\n'.join(current_paragraphs)))
                    current_paragraphs = []
                
                # Process the table
                table = Table(element, doc)
                table_text = process_table(table)
                if table_text:
                    content_parts.append(('table', table_text))
        
        # Add any remaining paragraphs
        if current_paragraphs:
            content_parts.append(('text', '\n'.join(current_paragraphs)))
        
        # print("\nProcessing headers and footers...")
        # Process headers and footers
        for section_idx, section in enumerate(doc.sections):
            # print(f"Processing section {section_idx + 1}")
            for header in section.header.paragraphs:
                if header.text.strip():
                    content_parts.append(('header', f"Header: {header.text}"))
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    content_parts.append(('footer', f"Footer: {footer.text}"))
        
        # print("\nDocument processing completed")
        return content_parts
    except Exception as e:
        # print(f"\nError during document processing: {str(e)}")
        raise Exception(f"Error extracting text from document: {str(e)}")

def preprocess_doc(file = {}, chunk_size: int = 1000, overlap: int = 200, blob_metadata = None) -> Dict[str, List[str]]:
    """
    Main preprocessing function that handles the entire document preprocessing pipeline.
    """
    print(f"\nStarting document preprocessing: {file["name"]}")
    
    if isinstance(file["content"], bytes):
        file["content"] = Document(io.BytesIO(file["content"]))
    print("\nStep 1: Extracting text from document")
    content_parts = extract_text_from_docx(file)
    print(f"Extracted {len(content_parts)} content parts from document")
    # print("Content parts preview:", content_parts)
    
    # Process each part according to its type
    processed_parts = []
    for content_type, content in content_parts:
        if content_type == 'table':
            # Keep tables as is
            processed_parts.append(content)
        else:
            # Clean non-table content
            processed_parts.append(clean_text(content))
    
    # Join all parts together
    final_text = '\n'.join(processed_parts)
    print(f"Extracted {len(final_text)} characters of processed text")
    print('Processed text preview:', final_text[:500])
    
    print("\nStep 2: Splitting into chunks")
    chunks = split_text_into_chunks(final_text, chunk_size, overlap)
    
    # Normalize chunks
    normalized_chunks = []
    for chunk in chunks:
        normalized_chunk = normalize_chunk(chunk)
        normalized_chunks.append(normalized_chunk)
        print(f"Normalized chunk: {normalized_chunk[:100]}...")
    
    print(f"Chunk length for document {file['name']} = {len(normalized_chunks)}")
    
    timestamp = datetime.fromtimestamp(time.time(), tz=timezone.utc)
    # Prepare metadata
    metadata = {
        # TODO: userId, clientId, timestamp
        'user_id': 'example_user',
        'client_id': 'example_client',
        'timestamp': timestamp.strftime('%Y-%m-%dT%H:%M:%S.%f')[:-3] + 'Z',
        'num_chunks': len(normalized_chunks),
        'total_chars': len(final_text)
    }
    metadata.update(blob_metadata or {})
    metadata["has_been_preprocessed"] = "True"
    
    print("\nPreprocessing completed successfully!")
    return {
        'result': {
            'chunks': normalized_chunks,
            'metadata': metadata
        }
    }

if __name__ == "__main__":
    try:
        start_time = time.time()
        doc_path = r"test_data/sample.docx"
        print("\n=== Starting Document Processing ===")
        
        with open(doc_path, 'rb') as f:
            file_content = f.read()
        
        file_dict = {"name": doc_path, "content": file_content}
        result = preprocess_doc(file_dict, blob_metadata=None)
        
        print("\n=== Processing Results ===")
        print(f"Processed document: {result['result']['metadata']['user_id']}")
        print(f"Number of chunks: {result['result']['metadata']['num_chunks']}")
        print(f"Total characters: {result['result']['metadata']['total_chars']}")
        print(f"Processing time: {time.time() - start_time:.2f} seconds")
        
        if result['result']['chunks']:
            print(f"\nFirst chunk example:\n{result['result']['chunks'][0]}\n")
            
    except Exception as e:
        print(f"\nError processing document: {str(e)}")
